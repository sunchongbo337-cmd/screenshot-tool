from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

paths = [
    r"C:\Users\lenovo\Desktop\中期报告.docx",
    r"C:\Users\lenovo\Desktop\中期报告-孙.docx",
]

for path in paths:
    print(f"\n=== {path} ===")
    p = Path(path)
    if not p.exists():
        print("FILE NOT FOUND")
        continue
    doc = Document(path)
    print("PARAGRAPHS:")
    paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paras.append((para.style.name if para.style else "", text))
    for i, (style, text) in enumerate(paras[:160], 1):
        print(f"{i:03d} [{style}] {text[:240]}")
    print("\nTABLES:")
    for ti, table in enumerate(doc.tables[:10], 1):
        print(f"TABLE {ti}")
        for row in table.rows[:20]:
            print(" | ".join(cell.text.strip() for cell in row.cells)[:500])
    print("\nCOMMENTS:")
    try:
        comments_part = doc.part.comments_part
        root = comments_part._element
        comments = []
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for comment in root.findall(f"{{{ns}}}comment"):
            cid = comment.get(f"{{{ns}}}id")
            author = comment.get(f"{{{ns}}}author")
            date = comment.get(f"{{{ns}}}date")
            texts = [t.text for t in comment.iter() if t.tag.endswith("}t") and t.text]
            body = "".join(texts).strip()
            print(f"- {cid} | {author} | {date}: {body[:400]}")
            comments.append((cid, author, date, body))
        if not comments:
            print("NO COMMENTS FOUND")
    except Exception as e:
        print(f"ERROR reading comments: {type(e).__name__}: {e}")
